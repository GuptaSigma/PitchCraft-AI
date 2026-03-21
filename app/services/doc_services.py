"""
DOC Export Service - PPTX layout parity
======================================

Generates Word-compatible documents (DOC) using the same theme system and
layout method names as pptx_service.py. Rendering relies on python-docx.
"""

from io import BytesIO
import json
import logging
import traceback
from typing import Any, Dict, List, Tuple

import requests
from pathlib import Path
from urllib.parse import urlparse, urljoin
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

class DOCService:
    """Generate DOC documents with pptx-style layouts."""

    def __init__(self) -> None:
        self.themes = self._build_themes()
        self.current_theme = self.themes['dialogue']
        self.page_width_inches = 13.3333
        self.page_height_inches = 7.5
        self.page_margin_inches = 0.5

    # Theme helpers
    def _build_themes(self) -> Dict[str, Dict[str, Tuple[int, int, int]]]:
        return {
            'dialogue': {'name': 'Dialogue White', 'bg': (255, 255, 255), 'text': (15, 23, 42), 'accent': (99, 102, 241), 'card': (248, 250, 252)},
            'alien': {'name': 'Alien Dark', 'bg': (15, 23, 42), 'text': (241, 245, 249), 'accent': (34, 211, 238), 'card': (30, 41, 59)},
            'wine': {'name': 'Wine Elegance', 'bg': (88, 28, 60), 'text': (255, 222, 200), 'accent': (244, 114, 182), 'card': (45, 11, 30)},
            'snowball': {'name': 'Snowball Blue', 'bg': (224, 242, 254), 'text': (30, 58, 138), 'accent': (14, 165, 233), 'card': (186, 230, 253)},
            'petrol': {'name': 'Petrol Steel', 'bg': (71, 85, 105), 'text': (241, 245, 249), 'accent': (14, 165, 233), 'card': (51, 65, 85)},
            'piano': {'name': 'Piano Contrast', 'bg': (255, 255, 255), 'text': (0, 0, 0), 'accent': (0, 0, 0), 'card': (245, 245, 245)},
            'sunset': {'name': 'Sunset Orange', 'bg': (254, 252, 232), 'text': (120, 53, 15), 'accent': (249, 115, 22), 'card': (254, 243, 199)},
            'midnight': {'name': 'Midnight Purple', 'bg': (30, 41, 59), 'text': (226, 232, 240), 'accent': (168, 85, 247), 'card': (15, 23, 42)},
        }

    def _calculate_brightness(self, rgb: Tuple[int, int, int]) -> float:
        """Calculate perceived brightness of RGB color (0-255)"""
        r, g, b = rgb
        return (r * 299 + g * 587 + b * 114) / 1000

    def _get_readable_text_color(self, background_rgb: Tuple[int, int, int]) -> RGBColor:
        brightness = self._calculate_brightness(background_rgb)
        if brightness > 128:
            return RGBColor(15, 23, 42)  # Dark text
        else:
            return RGBColor(255, 255, 255) # Light text

    def _get_readable_text_rgb(self, rgb: Tuple[int, int, int]) -> Tuple[int, int, int]:
        return (0, 0, 0) if self._calculate_brightness(rgb) > 128 else (255, 255, 255)

    def _fit_card_font_size(self, text: str) -> int:
        length = len(text or "")
        if length > 420: return 8
        if length > 320: return 9
        if length > 240: return 10
        return 11

    def _parse_content(self, content: Any) -> List[str]:
        """Parse content into list of points (Improved Regex Cleaning)"""
        if isinstance(content, list):
            return [str(item).strip() for item in content if str(item).strip()]
        
        text = str(content or '')
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Remove bullet markers, numbers, and Markdown characters
        import re
        cleaned = []
        for line in lines:
            # 1. Strip common bullet/number markers: "• ", "- ", "* ", "1. ", "01. ", etc.
            line = re.sub(r'^[\s•\-*]*([0-9]{1,2}[\.\)]\s*)?', '', line)
            
            # 2. Strip leading colon and spaces (common in AI output like "1. : Content")
            line = re.sub(r'^[:\s\t]*', '', line)
            
            # 3. Strip leading double asterisks (Markdown bold)
            line = re.sub(r'^\**\s*', '', line)
            
            # 4. Strip trailing asterisks
            line = re.sub(r'\**\s*$', '', line)
            
            # 5. Final trim
            line = line.strip()
            
            if line:
                cleaned.append(line)
        
        return cleaned

    # Data helpers
    def _detect_theme(self, presentation_data: Any) -> str:
        theme_name = getattr(presentation_data, 'theme', None)
        if theme_name is None and hasattr(presentation_data, '__dict__'):
            theme_name = presentation_data.__dict__.get('theme')
        if theme_name is None and isinstance(presentation_data, dict):
            theme_name = presentation_data.get('theme')
        if not theme_name or str(theme_name).strip() == '':
            logger.warning("Theme missing, defaulting to dialogue")
            return 'dialogue'
        normalized = str(theme_name).lower().strip()
        if normalized not in self.themes:
            logger.warning("Theme %s invalid, defaulting to dialogue", normalized)
            return 'dialogue'
        return normalized

    def _extract_slides(self, presentation_data: Any) -> Tuple[str, List[Dict[str, Any]]]:
        title = getattr(presentation_data, 'title', None) or 'Presentation'
        content = getattr(presentation_data, 'content', presentation_data)
        if isinstance(content, str):
            try:
                content = json.loads(content)
            except Exception:
                content = {}
        slides = content.get('slides', []) if isinstance(content, dict) else []
        return title, slides if isinstance(slides, list) else []

    def _download_image(
        self,
        url: str,
        size: Tuple[int, int] = (800, 600),
        cover: bool = False,
    ) -> BytesIO | None:
        if not url:
            return None

        url_str = str(url).strip()
        try:
            parsed = urlparse(url_str)
        except Exception:
            parsed = None

        if parsed and parsed.scheme in ("http", "https"):
            return self._download_remote_image(url_str, size, cover)

        local_stream = self._load_local_image(parsed.path if parsed else url_str, size, cover)
        if local_stream:
            return local_stream

        if parsed and not parsed.scheme and url_str.startswith("//"):
            return self._download_remote_image(f"https:{url_str}", size, cover)

        if parsed and not parsed.scheme:
            try:
                from flask import request

                base = request.host_url if request else None
                if base:
                    absolute_url = urljoin(base, parsed.path.lstrip("/"))
                    return self._download_remote_image(absolute_url, size, cover)
            except Exception:
                pass

        logger.warning("Image download failed: Unsupported or missing image source '%s'", url_str)
        return None

    def _prepare_image(self, image: Image.Image, size: Tuple[int, int], cover: bool) -> Image.Image:
        if not size:
            return image

        if cover:
            return self._resize_cover(image, size)

        img_copy = image.copy()
        img_copy.thumbnail(size, Image.Resampling.LANCZOS)
        return img_copy

    def _download_remote_image(self, url: str, size: Tuple[int, int], cover: bool) -> BytesIO | None:
        try:
            response = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'}, stream=True)
            if response.status_code != 200:
                logger.warning("Remote image fetch failed with status %s for '%s'", response.status_code, url)
                return None

            image = Image.open(BytesIO(response.content))
            image = self._prepare_image(image, size, cover)
            buffer = BytesIO()
            image.convert('RGB').save(buffer, format='PNG')
            buffer.seek(0)
            return buffer
        except Exception as exc:
            logger.warning("Image download failed: %s", exc)
            return None

    def _load_local_image(self, path: str, size: Tuple[int, int], cover: bool) -> BytesIO | None:
        if not path:
            return None

        normalized = path.lstrip("/")
        try:
            root_dir = Path(__file__).resolve().parents[2]
            candidate = (root_dir / normalized).resolve()

            if not str(candidate).startswith(str(root_dir)) or not candidate.exists():
                return None

            with candidate.open('rb') as handle:
                data = handle.read()

            image = Image.open(BytesIO(data))
            image = self._prepare_image(image, size, cover)
            buffer = BytesIO()
            image.convert('RGB').save(buffer, format='PNG')
            buffer.seek(0)
            return buffer
        except Exception as exc:
            logger.warning("Local image load failed for '%s': %s", path, exc)
            return None

    def _resize_cover(self, image: Image.Image, size: Tuple[int, int]) -> Image.Image:
        target_width, target_height = size
        if not target_width or not target_height:
            return image

        img_width, img_height = image.size
        if not img_width or not img_height:
            return image

        target_ratio = target_width / target_height
        image_ratio = img_width / img_height

        if image_ratio > target_ratio:
            scaled_height = target_height
            scaled_width = int(round(scaled_height * image_ratio))
        else:
            scaled_width = target_width
            scaled_height = int(round(scaled_width / image_ratio))

        resized = image.resize((scaled_width, scaled_height), Image.Resampling.LANCZOS)

        left = max(0, (scaled_width - target_width) // 2)
        top = max(0, (scaled_height - target_height) // 2)
        right = left + target_width
        bottom = top + target_height

        return resized.crop((left, top, right, bottom))

    def _add_text_overlay_to_image(self, img_stream, title, content_list):
        """Burn text directly into the image to simulate overlay in Word"""
        try:
            img = Image.open(img_stream)
            img = img.convert("RGBA")
            width, height = img.size
            
            # Create transparent layer for text
            text_layer = Image.new('RGBA', img.size, (0, 0, 0, 0))
            draw = ImageDraw.Draw(text_layer)
            
            # 1. Darken the image for contrast
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(0.5).convert("RGBA")

            # Load Fonts
            try:
                # Common Windows path
                font_path = "C:\\Windows\\Fonts\\arial.ttf"
                title_font = ImageFont.truetype(font_path, 60)
                text_font = ImageFont.truetype(font_path, 32)
            except:
                title_font = ImageFont.load_default()
                text_font = ImageFont.load_default()
            
            # Helper for text size (Pillow 10+ compatibility)
            def get_text_size(text, font):
                try:
                    left, top, right, bottom = font.getbbox(text)
                    return right - left, bottom - top
                except:
                    # Fallback for older Pillow
                    tw, th = draw.textsize(text, font=font)
                    return tw, th

            # Draw Title (Centered)
            title_text = str(title).upper()
            tw, th = get_text_size(title_text, title_font)
            draw.text(((width - tw) / 2, height * 0.25), title_text, font=title_font, fill=(255, 255, 255, 255))
            
            # Draw Content (Centered wrapping)
            y_offset = height * 0.25 + th + 60
            full_text = " ".join(content_list)
            
            words = full_text.split()
            lines = []
            current_line = []
            for word in words:
                current_line.append(word)
                lw, lh = get_text_size(" ".join(current_line), font=text_font)
                if lw > width * 0.85:
                    current_line.pop()
                    lines.append(" ".join(current_line))
                    current_line = [word]
            lines.append(" ".join(current_line))
            
            for line in lines[:6]: # Limit to 6 lines to fit
                lw, lh = get_text_size(line, font=text_font)
                draw.text(((width - lw) / 2, y_offset), line, font=text_font, fill=(245, 245, 245, 255))
                y_offset += lh + 20
            
            # Merge layers
            final_img = Image.alpha_composite(img, text_layer)
            
            # Save back
            output = BytesIO()
            final_img.convert("RGB").save(output, format='PNG')
            output.seek(0)
            return output
        except Exception as e:
            logger.error(f"Error in image overlay: {e}")
            return img_stream

    # Document helpers
    def _set_page_background(self, doc, rgb):
        """Set page background color using XML manipulation"""
        try:
            # Enable background shape display in settings
            settings = doc.settings.element
            bg_shape = OxmlElement('w:displayBackgroundShape')
            settings.append(bg_shape)

            # Get document element
            root = doc.element
            
            # Create background element
            background = OxmlElement('w:background')
            background.set(qn('w:color'), '%02x%02x%02x' % rgb)
            
            # Insert at the beginning of document element
            root.insert(0, background)
        except Exception as e:
            logger.warning(f"Could not set page background: {e}")

    def _add_page_color(self, paragraph, rgb):
        """Add background color to paragraph"""
        try:
            p = paragraph._element
            pPr = p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '%02x%02x%02x' % rgb)
            pPr.append(shd)
        except Exception as e:
            logger.warning(f"Could not add paragraph background: {e}")

    def _add_table_cell_color(self, cell, rgb):
        """Add background color to table cell"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '%02x%02x%02x' % rgb)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.warning(f"Could not add cell background: {e}")

    def _add_border_to_cell(self, cell, color_rgb, width=1):
        """Add border to table cell"""
        try:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(width * 8))  # width in eighths of a point
                border.set(qn('w:color'), '%02x%02x%02x' % color_rgb)
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        except Exception as e:
            logger.warning(f"Could not add cell border: {e}")

    def _init_document(self) -> Document:
        doc = Document()
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor(*self.current_theme['text'])

        section = doc.sections[0]
        section.page_width = Inches(self.page_width_inches)
        section.page_height = Inches(self.page_height_inches)
        section.left_margin = Inches(self.page_margin_inches)
        section.right_margin = Inches(self.page_margin_inches)
        section.top_margin = Inches(self.page_margin_inches)
        section.bottom_margin = Inches(self.page_margin_inches)
        return doc

    def _add_text(self, doc: Document, text: str, italic: bool = False, center: bool = False, size: int = 11, color: Tuple[int, int, int] = None) -> None:
        """Add text directly to document"""
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(*(color or self.current_theme['text']))
            run.font.size = Pt(size)
            run.italic = italic

    # Public API
    def generate(self, presentation_data: Any) -> bytes:
        try:
            theme_key = self._detect_theme(presentation_data)
            self.current_theme = self.themes[theme_key]
            title, slides = self._extract_slides(presentation_data)

            doc = self._init_document()
            
            # Set page background color
            self._set_page_background(doc, self.current_theme['bg'])
            
            self._create_cover(doc, title, theme_key, len(slides))
            if slides:
                doc.add_page_break()

            total = len(slides)
            for index, slide in enumerate(slides, start=1):
                layout = str(slide.get('layout', 'standard')).lower()
                builder = self._resolve_layout(layout, index)
                builder(doc, slide, index, total)
                if index < total:
                    doc.add_page_break()

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
        except Exception as exc:
            logger.error("DOC generation failed: %s", exc)
            traceback.print_exc()
            raise

    # Cover
    def _create_cover(self, doc: Document, title: str, theme_key: str, slides: int) -> None:
        """DOC Version: Improved cover"""
        # Add space
        for _ in range(3):
            doc.add_paragraph()

        # Container Table for centered content
        table = doc.add_table(rows=1, cols=1)
        table.width = Inches(12.33)
        cell = table.cell(0, 0)
        self._add_table_cell_color(cell, self.current_theme['bg'])
        self._add_border_to_cell(cell, self.current_theme['accent'], width=2)

        self._add_heading_to_cell(cell, title.upper(), center=True)
        self._add_text_to_cell(cell, f"THEME: {theme_key.upper()} • {slides} SLIDES", center=True, italic=True)
        self._add_text_to_cell(cell, "Gamma Presentation Document", center=True, size=10)

        # Bottom space
        for _ in range(5):
            doc.add_paragraph()

    # Layout dispatcher
    def _resolve_layout(self, layout: str, slide_index: int):
        mapping = {
            'hero_overlay': self._create_hero_overlay_slide,
            'hero': self._create_hero_overlay_slide,
            'centered': self._create_centered_slide,
            'split_box': self._create_fixed_split_box_slide_slide5 if slide_index == 5 else self._create_fixed_split_box_slide,
            'three_col': self._create_fixed_three_cards_slide,
            'grid_4': self._create_fixed_four_grid_slide,
            'fixed_information': self._create_executive_summary_slide if slide_index == 7 else self._create_fixed_image_cards_slide,
            'roadmap': self._create_fixed_roadmap_clean_slide,
            'fixed_mission': self._create_image_overlay_slide,
        }
        return mapping.get(layout, self._create_standard_slide)

    # Layout implementations (names mirror pptx_service)
    def _add_heading_to_cell(self, cell, text, center=False, size=24, color=None):
        """Add a heading to a table cell"""
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = p.add_run(str(text))
        run.bold = True
        run.font.size = Pt(size)
        
        # If no color, use theme accent
        if not color:
            color = self.current_theme['accent']
            
        run.font.color.rgb = RGBColor(*color)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _add_text_to_cell(self, cell, text, center=False, italic=False, size=14, color=None):
        """Add regular text to a table cell with theme-aware default color"""
        p = cell.add_paragraph()
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        
        # If no color, use theme text color OR readable fallback for light themes
        if not color:
            color = self.current_theme['text']
            # If the main theme text color might be invisible on its own BG
            # (Happens if someone set white text on light bg in theme definition)
            if self._calculate_brightness(self.current_theme['bg']) > 128 and self._calculate_brightness(color) > 128:
                color = (15, 23, 42) # Dark fallback
        
        run.font.color.rgb = RGBColor(*color)
        run.italic = italic
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def _create_container_table(self, doc, rows=1, cols=1):
        """Create a table that fills the length"""
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # We try to force width by setting column widths
        total_width_in = 12.33
        col_width = total_width_in / cols
        for i in range(cols):
            table.columns[i].width = Inches(col_width)
        return table

    def _create_hero_overlay_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """DOC Version: Slide 1 (Hero) with image-based text burning"""
        # Try multiple image sources
        img_url = slide.get('image') or slide.get('bg_image') or slide.get('bg_url')
        title = slide.get('title', 'Presentation')
        # Robust content extraction for Slide 1
        content = slide.get('content') or slide.get('mission') or slide.get('agenda') or ""
        content_list = self._parse_content(content)

        # Container Table
        table = self._create_container_table(doc)
        cell = table.cell(0, 0)
        self._add_table_cell_color(cell, self.current_theme['bg'])

        if img_url:
            # For DOC, we burn text into image for overlay look
            stream = self._download_image(img_url, size=(1280, 720), cover=True)
            if stream:
                # BAKE TEXT INTO IMAGE
                processed_stream = self._add_text_overlay_to_image(stream, title, content_list)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(processed_stream, width=Inches(12.0), height=Inches(6.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self._add_heading_to_cell(cell, title.upper(), center=True, size=32)
                for point in content_list:
                    self._add_text_to_cell(cell, point, center=True)
        else:
            self._add_heading_to_cell(cell, title.upper(), center=True, size=32)
            for point in content_list:
                self._add_text_to_cell(cell, point, center=True)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Intro • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_centered_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """DOC Version: Slide 2 (Centered) with image-based text burning"""
        # Try multiple image sources
        img_url = slide.get('image') or slide.get('bg_image') or slide.get('bg_url')
        title = slide.get('title', 'Presentation')
        # Robust content extraction for Slide 2
        content = slide.get('content') or slide.get('mission') or slide.get('agenda') or ""
        content_list = self._parse_content(content)

        # Container Table
        table = self._create_container_table(doc)
        cell = table.cell(0, 0)
        self._add_table_cell_color(cell, self.current_theme['bg'])

        if img_url:
            # BAKE TEXT INTO IMAGE
            stream = self._download_image(img_url, size=(1280, 720), cover=True)
            if stream:
                processed_stream = self._add_text_overlay_to_image(stream, title, content_list)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(processed_stream, width= Inches(12.0), height=Inches(6.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                self._add_heading_to_cell(cell, title.upper(), center=True, size=32)
                for point in content_list:
                    self._add_text_to_cell(cell, point, center=True)
        else:
            self._add_heading_to_cell(cell, title.upper(), center=True, size=32)
            for point in content_list:
                self._add_text_to_cell(cell, point, center=True)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Overview • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_three_cards_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """3-column card layout"""
        title = slide.get('title') or f'Slide {index}'
        
        # Theme-driven background (no forced accent background)
        bg_rgb = self.current_theme['bg']
        title_rgb = self.current_theme['text']

        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, bg_rgb)

        # Header with Title
        header_table = self._create_container_table(doc)
        self._add_heading_to_cell(header_table.cell(0,0), title, center=True, size=28, color=title_rgb)
        doc.add_paragraph() # Spacer

        # Cards Table
        cards = (slide.get('columns') or [])[:3]
        if not cards:
            points = self._parse_content(slide.get('content', ''))
            cards = [{'title': f'Point {i+1}', 'content': p} for i, p in enumerate(points[:3])]

        table = self._create_container_table(doc, cols=len(cards) if cards else 3)
        for i, card in enumerate(cards):
            cell = table.cell(0, i)
            self._add_table_cell_color(cell, self.current_theme['card'])
            self._add_border_to_cell(cell, self.current_theme['accent'], width=1)
            
            # Badge
            badge_p = cell.add_paragraph()
            badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            badge_run = badge_p.add_run(f" {i+1} ")
            badge_run.bold = True
            badge_run.font.size = Pt(18)
            badge_run.font.color.rgb = self._get_readable_text_color(self.current_theme['accent'])
            pPr = badge_p._element.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '%02x%02x%02x' % self.current_theme['accent'])
            pPr.append(shd)

            self._add_text_to_cell(cell, card.get('title') or 'Feature', center=True, size=16, color=self.current_theme['accent'])
            body = card.get('content') or 'Detail content here.'
            self._add_text_to_cell(cell, body, size=self._fit_card_font_size(body))

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Columns • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_four_grid_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """2x2 Grid card layout"""
        title = slide.get('title') or f'Slide {index}'
        
        header_table = self._create_container_table(doc)
        self._add_heading_to_cell(header_table.cell(0,0), title, center=True, size=28)
        doc.add_paragraph()

        items = slide.get('items') or []
        if not items:
            points = self._parse_content(slide.get('content', ''))
            items = [{'title': f'Item {i+1}', 'content': p} for i, p in enumerate(points[:4])]
        
        table = self._create_container_table(doc, rows=2, cols=2)
        for i in range(4):
            row_idx, col_idx = divmod(i, 2)
            cell = table.cell(row_idx, col_idx)
            self._add_table_cell_color(cell, self.current_theme['card'])
            self._add_border_to_cell(cell, self.current_theme['accent'], width=1)
            
            if i < len(items):
                item = items[i]
                self._add_text_to_cell(cell, item.get('title') or f'Point {i+1}', size=16, color=self.current_theme['accent'])
                body = item.get('content') or 'Information goes here.'
                self._add_text_to_cell(cell, body, size=self._fit_card_font_size(body))

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Grid • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_image_cards_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Gallery style: Image + Cards side-by-side"""
        title = slide.get('title') or f'Slide {index}'
        
        table = self._create_container_table(doc, cols=2)
        # Left: Image
        left_cell = table.cell(0, 0)
        img_url = slide.get('image') or slide.get('bg_image') or slide.get('bg_url')
        if img_url:
            stream = self._download_image(img_url, size=(800, 800))
            if stream:
                p = left_cell.paragraphs[0]
                p.add_run().add_picture(stream, width=Inches(5.0))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Right: Cards
        right_cell = table.cell(0, 1)
        self._add_heading_to_cell(right_cell, title, size=22)
        cards = slide.get('cards') or [{'title': 'Details', 'content': slide.get('content')}]
        for card in cards[:3]:
            self._add_text_to_cell(right_cell, card.get('title') or 'Info', italic=True, size=14, color=self.current_theme['accent'])
            self._add_text_to_cell(right_cell, card.get('content') or 'Description.', size=11)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Gallery • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_executive_summary_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Modern executive summary with design elements (Slide 7 special)"""
        # White background for summary
        self._set_page_background(doc, (255, 255, 255))
        
        table = self._create_container_table(doc, cols=2)
        table.columns[0].width = Inches(8.5)
        table.columns[1].width = Inches(2.5)
        
        # Left: Main title
        left_cell = table.cell(0, 0)
        title = slide.get('title', 'Executive Summary')
        self._add_heading_to_cell(left_cell, title, size=36, color=(30, 41, 59))
        
        # Right: Badge
        right_cell = table.cell(0, 1)
        self._add_table_cell_color(right_cell, (30, 58, 138))
        badge_para = right_cell.paragraphs[0]
        badge_run = badge_para.add_run("EXECUTIVE\nSUMMARY")
        badge_run.font.size = Pt(9)
        badge_run.font.bold = True
        badge_run.font.color.rgb = RGBColor(255, 255, 255)
        badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Accent line
        doc.add_paragraph()
        accent_para = doc.add_paragraph()
        accent_run = accent_para.add_run(" " * 150)
        accent_run.font.size = Pt(3)
        pPr = accent_para._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '%02x%02x%02x' % self.current_theme['accent'])
        pPr.append(shd)

        # Content sections
        content_list = self._parse_content(slide.get('content', ''))
        while len(content_list) < 3:
            content_list.append(f"Section {len(content_list) + 1} summary...")

        for point in content_list[:3]:
            # Section Label
            label_p = doc.add_paragraph()
            label_r = label_p.add_run("KEY HIGHLIGHT")
            label_r.font.size = Pt(10)
            label_r.font.bold = True
            label_r.font.color.rgb = RGBColor(100, 116, 139)
            
            # Content
            self._add_text(doc, point, size=13, color=(51, 65, 85))
            doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Summary • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True
        # Revert page background for next slides if needed (usually handled by next slide builder)

    def _create_fixed_roadmap_clean_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Vertical roadmap implementation matching docx_services.py"""
        table = self._create_container_table(doc)
        cell = table.cell(0, 0)
        self._add_table_cell_color(cell, self.current_theme['bg'])
        
        self._add_heading_to_cell(cell, slide.get('title', f'Slide {index}'), center=True, size=32)
        doc.add_paragraph()
        
        content = self._parse_content(slide.get('content', ''))
        steps = content[:5]
        if not steps:
            steps = ["Phase 1: Planning", "Phase 2: Execution", "Phase 3: Launch"]
            
        # Roadmap table (nested for layout control)
        roadmap_table = doc.add_table(rows=len(steps), cols=2)
        roadmap_table.columns[0].width = Inches(1.5)
        roadmap_table.columns[1].width = Inches(10)
        
        circle_text_color = self._get_readable_text_color(self.current_theme['accent'])
        
        for i, step in enumerate(steps):
            row = roadmap_table.rows[i]
            
            # Left: Number Badge
            num_cell = row.cells[0]
            self._add_table_cell_color(num_cell, self.current_theme['accent'])
            p = num_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f" {i+1} ")
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = circle_text_color
            
            # Right: Content
            content_cell = row.cells[1]
            self._add_table_cell_color(content_cell, self.current_theme['card'])
            self._add_text_to_cell(content_cell, step, size=16)
            
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Project Roadmap • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_split_box_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Split layout: Two columns of metrics/points"""
        title = slide.get('title') or f'Slide {index}'
        header_table = self._create_container_table(doc)
        self._add_heading_to_cell(header_table.cell(0,0), title, center=True, size=28)
        doc.add_paragraph()

        table = self._create_container_table(doc, cols=2)
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        self._add_table_cell_color(left_cell, self.current_theme['card'])
        self._add_table_cell_color(right_cell, self.current_theme['card'])
        self._add_border_to_cell(left_cell, self.current_theme['accent'])
        self._add_border_to_cell(right_cell, self.current_theme['accent'])

        left_points = self._parse_content(slide.get('content') or 'Highlights')
        right_content = slide.get('secondary_content') or slide.get('notes') or 'Details'
        right_points = self._parse_content(right_content)

        for p in left_points[:6]:
            self._add_text_to_cell(left_cell, f"• {p}", size=11)
        for p in right_points[:6]:
            self._add_text_to_cell(right_cell, f"• {p}", size=11)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Metrics • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_split_box_slide_slide5(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Special Split Box for Slide 5: Image left, 2+1 grid dark panel right"""
        table = self._create_container_table(doc, cols=2)
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(6.3)
        
        # LEFT: Image
        left_cell = table.cell(0, 0)
        img_url = slide.get('image')
        if img_url:
            stream = self._download_image(img_url, size=(700, 900), cover=True)
            if stream:
                p = left_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(stream, width=Inches(4.5))
        
        # RIGHT: Color Panel
        right_cell = table.cell(0, 1)
        is_light = self._calculate_brightness(self.current_theme['bg']) > 128
        panel_rgb = self.current_theme['accent'] if is_light else self.current_theme['bg']
        text_color = (255, 255, 255) if is_light else self.current_theme['text']
        accent_color = (255, 255, 255) if is_light else self.current_theme['accent']
            
        self._add_table_cell_color(right_cell, panel_rgb)
        
        title = slide.get('title', f'Slide {index}')
        self._add_heading_to_cell(right_cell, title.upper(), size=24, color=text_color)
        
        content = self._parse_content(slide.get('content', ''))
        
        # Grid layout for content cards
        card_table = right_cell.add_table(rows=2, cols=2)
        card_points = content[:3]
        while len(card_points) < 3: card_points.append("Innovative solution detail for this section.")

        for i, point in enumerate(card_points):
            r = 0 if i < 2 else 1
            c = i if i < 2 else 0
            cell = card_table.cell(r, c)
            
            # Badge 01, 02, 03
            num_p = cell.add_paragraph()
            num_run = num_p.add_run(f"0{i+1}")
            num_run.bold = True
            num_run.font.size = Pt(18)
            num_run.font.color.rgb = RGBColor(*accent_color)
            
            # Point text
            text_p = cell.add_paragraph()
            text_run = text_p.add_run(point)
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = RGBColor(*text_color)
            text_p.paragraph_format.space_after = Pt(10)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Strategy • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_image_overlay_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """DOC Version: Image Overlay layout with 30/70 split parity"""
        img_url = slide.get('image') or slide.get('bg_image') or slide.get('bg_url')
        title = slide.get('title', f'Slide {index}')
        content = slide.get('content', '') or slide.get('mission', '') or slide.get('agenda', '')
        content_list = self._parse_content(content)
        
        is_light = self._calculate_brightness(self.current_theme['bg']) > 128

        # Container Table (30/70 split)
        table = self._create_container_table(doc, cols=2)
        table.columns[0].width = Inches(4.0)  # Left (30%)
        table.columns[1].width = Inches(8.3)  # Right (70%)
        
        # LEFT: Content Card
        left_cell = table.cell(0, 0)
        box_color = self.current_theme['accent'] if is_light else self.current_theme['bg']
        text_color = (255, 255, 255) if is_light else self.current_theme['text']
        
        self._add_table_cell_color(left_cell, box_color)
        
        # "KEY INSIGHT" Label
        label_p = left_cell.add_paragraph()
        run = label_p.add_run("KEY INSIGHT")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*text_color)
        
        # Title
        self._add_heading_to_cell(left_cell, title.upper(), size=22, color=text_color)
        
        # Points
        for point in content_list[:4]:
            p = left_cell.add_paragraph()
            p_run = p.add_run(f"• {point}")
            p_run.font.size = Pt(11)
            p_run.font.color.rgb = RGBColor(*text_color)
            p.paragraph_format.space_after = Pt(6)

        # RIGHT: Large Image
        right_cell = table.cell(0, 1)
        if img_url:
            stream = self._download_image(img_url, size=(1600, 900), cover=True)
            if stream:
                paragraph = right_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(stream, width=Inches(8.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self._add_table_cell_color(right_cell, (241, 245, 249))

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Vision • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True

    def _create_fixed_mission_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        self._create_image_overlay_slide(doc, slide, index, total)

    def _create_standard_slide(self, doc: Document, slide: Dict[str, Any], index: int, total: int) -> None:
        """Standard presentation slide layout"""
        title = slide.get('title') or f'Slide {index}'
        
        table = self._create_container_table(doc)
        cell = table.cell(0, 0)
        self._add_table_cell_color(cell, self.current_theme['bg'])
        
        self._add_heading_to_cell(cell, title, size=28)
        
        subtitle = slide.get('subtitle')
        if subtitle:
            self._add_text_to_cell(cell, str(subtitle), italic=True, size=16, color=self.current_theme['accent'])
        
        # Robust content extraction (mission/agenda)
        raw_content = slide.get('content') or slide.get('mission') or slide.get('agenda') or 'No details provided.'
        points = self._parse_content(raw_content)
        for point in points:
            self._add_text_to_cell(cell, f"• {point}", size=14)
            
        img_url = slide.get('image') or slide.get('bg_image') or slide.get('bg_url')
        if img_url:
            stream = self._download_image(img_url, size=(1000, 500))
            if stream:
                p = cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(stream, width=Inches(8.0))

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run(f"Details • Slide {index} of {total}")
        run.font.size = Pt(8)
        run.italic = True
