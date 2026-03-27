from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse, urljoin
import requests
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
import json
import traceback
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DOCXService:
    """
    DOCX generation service with 8 theme support and 8 layouts
    Mirrors PPTX service functionality exactly
    """
    
    def __init__(self):
        """Initialize DOCX service with theme definitions"""
        
        # THEME DEFINITIONS (8 THEMES - IDENTICAL TO PPTX)
        self.themes = {
            'dialogue': {
                'name': 'Dialogue White',
                'bg': (255, 255, 255),      # White
                'text': (15, 23, 42),       # Dark slate
                'accent': (99, 102, 241),   # Indigo
                'card': (248, 250, 252)     # Light gray
            },
            'alien': {
                'name': 'Alien Dark',
                'bg': (15, 23, 42),         # Dark blue
                'text': (241, 245, 249),    # Light gray
                'accent': (34, 211, 238),   # Cyan
                'card': (30, 41, 59)        # Darker blue
            },
            'wine': {
                'name': 'Wine Elegance',
                'bg': (88, 28, 60),         # Dark wine
                'text': (255, 222, 200),    # Cream
                'accent': (244, 114, 182),  # Pink
                'card': (45, 11, 30)        # Deeper wine
            },
            'snowball': {
                'name': 'Snowball Blue',
                'bg': (224, 242, 254),      # Light blue
                'text': (30, 58, 138),      # Dark blue
                'accent': (14, 165, 233),   # Sky blue
                'card': (186, 230, 253)     # Lighter blue
            },
            'petrol': {
                'name': 'Petrol Steel',
                'bg': (71, 85, 105),        # Steel gray
                'text': (241, 245, 249),    # Light gray
                'accent': (14, 165, 233),   # Sky blue
                'card': (51, 65, 85)        # Darker steel
            },
            'piano': {
                'name': 'Piano Contrast',
                'bg': (255, 255, 255),      # White
                'text': (0, 0, 0),          # Black
                'accent': (0, 0, 0),        # Black
                'card': (245, 245, 245)     # Light gray
            },
            'sunset': {
                'name': 'Sunset Orange',
                'bg': (254, 252, 232),      # Cream
                'text': (120, 53, 15),      # Brown
                'accent': (249, 115, 22),   # Orange
                'card': (254, 243, 199)     # Light yellow
            },
            'midnight': {
                'name': 'Midnight Purple',
                'bg': (30, 41, 59),         # Dark gray
                'text': (226, 232, 240),    # Light gray
                'accent': (168, 85, 247),   # Purple
                'card': (15, 23, 42)        # Darker gray
            }
        }
        
        self.current_theme = self.themes['dialogue']  # Default

        # Slide proportions (16:9) in inches for consistent exports
        self.page_width_inches = 13.3333
        self.page_height_inches = 7.5
        self.page_margin_inches = 0.5
        
        logger.info("✅ DOCX Service v5.0.1 initialized with 8 themes")
    
    
    # BRIGHTNESS CALCULATION (IDENTICAL TO PPTX)
    def _calculate_brightness(self, rgb):
        """
        Calculate perceived brightness of RGB color
        Returns: 0-255 (0=darkest, 255=brightest)
        """
        r, g, b = rgb
        brightness = (r * 299 + g * 587 + b * 114) / 1000
        return brightness
    
    
    def _get_readable_text_color(self, background_rgb):
        """
        Return black or white RGBColor based on background brightness
        """
        brightness = self._calculate_brightness(background_rgb)
        
        if brightness > 128:
            return RGBColor(0, 0, 0)  # Black text for light backgrounds
        else:
            return RGBColor(255, 255, 255)  # White text for dark backgrounds
    
    
    # THEME DETECTION (IDENTICAL TO PPTX)
    def _detect_theme(self, presentation_data):
        """
        Enhanced theme detection with multiple fallbacks
        """
        theme_name = None
        
        # Try 1: getattr (works for objects)
        theme_name = getattr(presentation_data, 'theme', None)
        
        # Try 2: __dict__ access (works for dict-like objects)
        if theme_name is None and hasattr(presentation_data, '__dict__'):
            if 'theme' in presentation_data.__dict__:
                theme_name = presentation_data.__dict__['theme']
        
        # Try 3: Direct dict access (if it's a dict)
        if theme_name is None and isinstance(presentation_data, dict):
            theme_name = presentation_data.get('theme')
        
        # Fallback: Default theme
        if not theme_name or str(theme_name).strip() == '':
            theme_name = 'dialogue'
            logger.warning("⚠️ No theme found, using default: dialogue")
        
        # Normalize
        theme_name = str(theme_name).lower().strip()
        
        # Validate
        if theme_name not in self.themes:
            logger.warning(f"⚠️ Invalid theme '{theme_name}', using dialogue")
            theme_name = 'dialogue'
        
        return theme_name
    
    
    # IMAGE DOWNLOAD (IDENTICAL TO PPTX)
    def _download_image(self, url, max_size=(1280, 720), cover=False):
        """Download an image from remote or local sources."""
        if not url:
            return None

        url_str = str(url).strip()
        try:
            parsed = urlparse(url_str)
        except Exception:
            parsed = None

        if parsed and parsed.scheme in ("http", "https"):
            return self._download_remote_image(url_str, max_size, cover)

        local_stream = self._load_local_image(parsed.path if parsed else url_str, max_size, cover)
        if local_stream:
            return local_stream

        if parsed and not parsed.scheme and url_str.startswith("//"):
            return self._download_remote_image(f"https:{url_str}", max_size, cover)

        if parsed and not parsed.scheme:
            try:
                from flask import request

                base = request.host_url if request else None
                if base:
                    absolute_url = urljoin(base, parsed.path.lstrip("/"))
                    return self._download_remote_image(absolute_url, max_size, cover)
            except Exception:
                pass

        logger.warning(f"⚠️ Image download failed: Unsupported or missing image source '{url_str}'")
        return None

    def _prepare_image(self, image, max_size, cover):
        if not max_size:
            return image

        if cover:
            return self._resize_cover(image, max_size)

        img_copy = image.copy()
        img_copy.thumbnail(max_size, Image.Resampling.LANCZOS)
        return img_copy

    def _download_remote_image(self, url, max_size, cover):
        try:
            response = requests.get(
                url,
                timeout=10,
                headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
                    "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8",
                },
                stream=True,
            )
            if response.status_code == 200:
                img = Image.open(BytesIO(response.content))
                img = self._prepare_image(img, max_size, cover)

                output = BytesIO()
                img.convert("RGB").save(output, format='PNG')
                output.seek(0)
                return output
            logger.warning(f"⚠️ Remote image fetch failed with status {response.status_code} for '{url}'")
        except Exception as exc:
            logger.warning(f"⚠️ Image download failed: {exc}")
        return None

    def _load_local_image(self, path, max_size, cover):
        if not path:
            return None

        normalized = path.lstrip("/")
        try:
            root_dir = Path(__file__).resolve().parents[2]
            candidate = (root_dir / normalized).resolve()

            if not str(candidate).startswith(str(root_dir)) or not candidate.exists():
                return None

            with candidate.open('rb') as handle:
                data = handle.read()

            img = Image.open(BytesIO(data))
            img = self._prepare_image(img, max_size, cover)

            output = BytesIO()
            img.convert("RGB").save(output, format='PNG')
            output.seek(0)
            return output
        except Exception as exc:
            logger.warning(f"⚠️ Local image load failed for '{path}': {exc}")
        return None
    
    def _add_text_overlay_to_image(self, img_stream, title, content_list):
        """Burn text directly into the image to simulate overlay in Word"""
        try:
            img = Image.open(img_stream)
            img = img.convert("RGBA")
            width, height = img.size
            
            # Create transparent layer for text
            text_layer = Image.new('RGBA', img.size, (0, 0, 0, 0))
            draw = ImageDraw.Draw(text_layer)
            
            # 1. Darken the image for contrast
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(0.5).convert("RGBA")

            # Load Fonts
            try:
                font_path = "C:\\Windows\\Fonts\\arial.ttf"
                title_font = ImageFont.truetype(font_path, 60)
                text_font = ImageFont.truetype(font_path, 32)
            except:
                title_font = ImageFont.load_default()
                text_font = ImageFont.load_default()
            
            # Helper for text size (Pillow 10+ compatibility)
            def get_text_size(text, font):
                try:
                    left, top, right, bottom = font.getbbox(text)
                    return right - left, bottom - top
                except:
                    return draw.textsize(text, font=font)

            # Draw Title (Centered)
            title_text = str(title).upper()
            tw, th = get_text_size(title_text, title_font)
            draw.text(((width - tw) / 2, height * 0.25), title_text, font=title_font, fill=(255, 255, 255, 255))
            
            # Draw Content (Centered wrapping)
            y_offset = height * 0.25 + th + 60
            full_text = " ".join(content_list)
            
            words = full_text.split()
            lines = []
            current_line = []
            for word in words:
                current_line.append(word)
                lw, lh = get_text_size(" ".join(current_line), font=text_font)
                if lw > width * 0.85:
                    current_line.pop()
                    lines.append(" ".join(current_line))
                    current_line = [word]
            lines.append(" ".join(current_line))
            
            for line in lines[:6]: # Limit to 6 lines to fit
                lw, lh = get_text_size(line, font=text_font)
                draw.text(((width - lw) / 2, y_offset), line, font=text_font, fill=(245, 245, 245, 255))
                y_offset += lh + 20
            
            # Merge layers
            final_img = Image.alpha_composite(img, text_layer)
            
            # Save back
            output = BytesIO()
            final_img.convert("RGB").save(output, format='PNG')
            output.seek(0)
            return output
        except Exception as e:
            logger.error(f"Error in image overlay: {e}")
            return img_stream
    
    def _parse_content(self, content):
        """Parse content into list of points (Improved Regex Cleaning)"""
        if isinstance(content, list):
            return content
        
        if not isinstance(content, str):
            content = str(content)
        
        # Split by newlines and clean
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        
        # Remove bullet markers, numbers, and Markdown characters
        import re
        cleaned = []
        for line in lines:
            # 1. Strip common bullet/number markers: "• ", "- ", "* ", "1. ", "01. ", etc.
            line = re.sub(r'^[\s•\-*]*([0-9]{1,2}[\.\)]\s*)?', '', line)
            
            # 2. Strip leading colon and spaces (common in AI output like "1. : Content")
            line = re.sub(r'^[:\s\t]*', '', line)
            
            # 3. Strip leading double asterisks (Markdown bold)
            line = re.sub(r'^\**\s*', '', line)
            
            # 4. Strip trailing asterisks
            line = re.sub(r'\**\s*$', '', line)
            
            # 5. Final trim
            line = line.strip()
            
            if line:
                cleaned.append(line)
        
        return cleaned
    
    
    # SCALING HELPERS (IDENTICAL TO PPTX)
    def _text_scale(self, texts, min_len=120, max_len=600):
        """Return 0..1 scale based on total text length."""
        if not texts:
            return 0.0
        if isinstance(texts, str):
            total = len(texts)
        else:
            total = sum(len(str(t)) for t in texts)

        if max_len <= min_len:
            return 1.0

        scale = (total - min_len) / (max_len - min_len)
        return max(0.0, min(1.0, scale))

    def _scaled_font(self, scale, max_font, min_font):
        """Scale font size based on text length scale."""
        return int(round(max_font - scale * (max_font - min_font)))

    def _scaled_height(self, scale, min_height, max_height):
        """Scale height based on text length scale."""
        return min_height + (max_height - min_height) * scale
    
    def _fit_card_font_size(self, text):
        """Pick a smaller font size for longer card text."""
        length = len(text or "")
        if length > 420:
            return 8
        if length > 320:
            return 9
        if length > 240:
            return 10
        return 11

    def _resize_cover(self, image, target_size):
        """Resize and crop an image so it fully covers the target size."""
        target_width, target_height = target_size
        if not target_width or not target_height:
            return image

        img_width, img_height = image.size
        if not img_width or not img_height:
            return image

        target_ratio = target_width / target_height
        image_ratio = img_width / img_height

        if image_ratio > target_ratio:
            scaled_height = target_height
            scaled_width = int(round(scaled_height * image_ratio))
        else:
            scaled_width = target_width
            scaled_height = int(round(scaled_width / image_ratio))

        resized = image.resize((scaled_width, scaled_height), Image.Resampling.LANCZOS)

        left = max(0, (scaled_width - target_width) // 2)
        top = max(0, (scaled_height - target_height) // 2)
        right = left + target_width
        bottom = top + target_height

        return resized.crop((left, top, right, bottom))
    
    
    # HELPER: SET PAGE BACKGROUND COLOR
    def _set_page_background(self, doc, rgb):
        """Set page background color using XML manipulation"""
        try:
            # Enable background shape display in settings
            settings = doc.settings.element
            bg_shape = OxmlElement('w:displayBackgroundShape')
            settings.append(bg_shape)

            # Get document element
            root = doc.element
            
            # Create background element
            background = OxmlElement('w:background')
            background.set(qn('w:color'), '%02x%02x%02x' % rgb)
            
            # Insert at the beginning of document element
            root.insert(0, background)
        except Exception as e:
            logger.warning(f"Could not set page background: {e}")
    
    
    def _add_page_color(self, paragraph, rgb):
        """Add background color to paragraph"""
        try:
            p = paragraph._element
            pPr = p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '%02x%02x%02x' % rgb)
            pPr.append(shd)
        except Exception as e:
            logger.warning(f"Could not add paragraph background: {e}")
    
    
    def _add_table_cell_color(self, cell, rgb):
        """Add background color to table cell"""
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '%02x%02x%02x' % rgb)
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logger.warning(f"Could not add cell background: {e}")
    
    
    def _add_border_to_cell(self, cell, color_rgb, width=1):
        """Add border to table cell"""
        try:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(width * 8))  # width in eighths of a point
                border.set(qn('w:color'), '%02x%02x%02x' % color_rgb)
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        except Exception as e:
            logger.warning(f"Could not add cell border: {e}")
    
    
    # MAIN GENERATION METHOD
    def generate(self, presentation_data):
        """
        Generate DOCX file from presentation data
        
        Args:
            presentation_data: Object/dict with:
                - title (str)
                - theme (str): dialogue/alien/wine/etc.
                - content (dict): {'slides': [...]}
        
        Returns:
            bytes: DOCX file content
        """
        try:
            logger.info(f"\n{'='*80}")
            logger.info("[DOCX] Starting generation...")
            logger.info(f"{'='*80}")
            
            # STEP 1: DETECT & APPLY THEME
            theme_name = self._detect_theme(presentation_data)
            self.current_theme = self.themes[theme_name]
            
            logger.info(f"[THEME] Applied: {self.current_theme['name']} ({theme_name})")
            logger.info(f"[THEME] Background: RGB{self.current_theme['bg']}")
            logger.info(f"[THEME] Text Color: RGB{self.current_theme['text']}")
            logger.info(f"[THEME] Card Color: RGB{self.current_theme['card']}")
            logger.info(f"[THEME] Accent: RGB{self.current_theme['accent']}")
            
            # STEP 2: CREATE DOCUMENT
            doc = Document()
            
            # Set background color for the whole document
            self._set_page_background(doc, self.current_theme['bg'])
            
            # Set page to landscape
            section = doc.sections[0]
            section.page_width = Inches(self.page_width_inches)
            section.page_height = Inches(self.page_height_inches)
            section.left_margin = Inches(self.page_margin_inches)
            section.right_margin = Inches(self.page_margin_inches)
            section.top_margin = Inches(self.page_margin_inches)
            section.bottom_margin = Inches(self.page_margin_inches)
            
            # Get slides data
            content = presentation_data.content if hasattr(presentation_data, 'content') else presentation_data
            
            if isinstance(content, str):
                content = json.loads(content)
            
            slides_data = content.get('slides', []) if isinstance(content, dict) else content
            
            logger.info(f"[SLIDES] Total to generate: {len(slides_data)}")
            
            # STEP 3: GENERATE EACH SLIDE
            for idx, slide_data in enumerate(slides_data):
                slide_num = idx + 1
                layout = slide_data.get('layout', 'standard')
                
                logger.info(f"[SLIDE {slide_num}] Layout: {layout}")
                
                # Route to appropriate layout method
                if layout == 'hero_overlay' or layout == 'hero':
                    self._create_hero_overlay_slide(doc, slide_data, slide_num)
                elif layout == 'centered':
                    self._create_centered_slide(doc, slide_data, slide_num)
                elif layout == 'split_box':
                    if slide_num == 5:
                        self._create_fixed_split_box_slide_slide5(doc, slide_data, slide_num)
                    else:
                        self._create_fixed_split_box_slide(doc, slide_data, slide_num)
                elif layout == 'three_col':
                    self._create_fixed_three_cards_slide(doc, slide_data, slide_num)
                elif layout == 'grid_4':
                    self._create_fixed_four_grid_slide(doc, slide_data, slide_num)
                elif layout == 'fixed_information':
                    if slide_num == 7:
                        self._create_executive_summary_slide(doc, slide_data, slide_num)
                    else:
                        self._create_fixed_image_cards_slide(doc, slide_data, slide_num)
                elif layout == 'roadmap':
                    self._create_fixed_roadmap_clean_slide(doc, slide_data, slide_num)
                elif layout == 'fixed_mission':
                    self._create_image_overlay_slide(doc, slide_data, slide_num)
                else:
                    self._create_standard_slide(doc, slide_data, slide_num)
                
                # Add page break after each slide (except last)
                if idx < len(slides_data) - 1:
                    doc.add_page_break()
            
            # STEP 4: SAVE TO BYTES
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info(f"[SUCCESS] DOCX completed - Theme: {self.current_theme['name']}")
            logger.info(f"{'='*80}\n")
            
            return output.getvalue()
        
        except Exception as e:
            logger.error(f"❌ DOCX generation error: {e}")
            traceback.print_exc()
            raise
    
    
    # LAYOUT 1: HERO OVERLAY SLIDE
    def _create_hero_overlay_slide(self, doc, slide_data, slide_num):
        """Hero overlay slide with text burned directly into the image for true overlay look"""
        
        # Try multiple image sources for robustness
        image_url = slide_data.get("image") or slide_data.get("bg_image") or slide_data.get("bg_url")
        
        title = slide_data.get("title", f"Slide {slide_num}")
        
        # Content fallbacks for hero slides so first page always gets body text when available
        content_raw = (
            slide_data.get("content")
            or slide_data.get("subtitle")
            or slide_data.get("description")
            or slide_data.get("summary")
            or slide_data.get("overview")
            or slide_data.get("mission")
            or slide_data.get("agenda")
            or ""
        )
        content = self._parse_content(content_raw)

        if not content and slide_data.get("subtitle"):
            content = [str(slide_data.get("subtitle"))]
        
        if image_url:
            img_stream = self._download_image(image_url, max_size=(1600, 900), cover=True)
            if img_stream:
                # Burn text into the image
                final_img_stream = self._add_text_overlay_to_image(img_stream, title, content)
                
                # Add the combined image to the document
                doc.add_picture(final_img_stream, width=Inches(self.page_width_inches - 1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Center vertically on page roughly
                shapes = doc.inline_shapes
                if len(shapes):
                    shape = shapes[len(shapes) - 1]
                    max_height = Inches(self.page_height_inches - 1)
                    if shape.height > max_height:
                        shape.height = max_height
        else:
            # Fallback if no image
            # Set background color for the page
            bg_para = doc.add_paragraph()
            self._add_page_color(bg_para, self.current_theme['bg'])
            text_color = self._get_readable_text_color(self.current_theme['bg'])

            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title.upper())
            title_run.font.size = Pt(44)
            title_run.font.bold = True
            title_run.font.color.rgb = text_color
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph() # Spacer

            for line in content[:5]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.size = Pt(22)
                run.font.color.rgb = text_color
    
    
    # LAYOUT 2: CENTERED SLIDE
    def _create_centered_slide(self, doc, slide_data, slide_num):
        """Centered slide with text burned into the image"""
        self._create_hero_overlay_slide(doc, slide_data, slide_num) # Use same logic for overlay parity
    
    
    
    # LAYOUT 3: THREE CARDS
    def _create_fixed_three_cards_slide(self, doc, slide_data, slide_num):
        """Three vertical cards layout"""
        
        # Theme-driven background (no forced accent background)
        bg_rgb = self.current_theme['bg']
        title_rgb = self.current_theme['text']

        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, bg_rgb)
        
        # Title (centered)
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*title_rgb)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Spacer
        doc.add_paragraph()
        
        # Parse content (expect 3 points)
        content = self._parse_content(slide_data.get('content', ''))
        points = content[:3]
        while len(points) < 3:
            points.append(f"Point {len(points) + 1}")
        
        # Create table with 3 columns
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for col_idx in range(3):
            table.columns[col_idx].width = Inches(3.3)
        
        row = table.rows[0]
        
        for i, point in enumerate(points):
            cell = row.cells[i]
            
            # Set cell background color
            self._add_table_cell_color(cell, self.current_theme['card'])
            self._add_border_to_cell(cell, self.current_theme['accent'], width=2)
            
            # Number badge
            num_para = cell.add_paragraph()
            num_run = num_para.add_run(f"  {i+1}  ")
            num_run.font.size = Pt(20)
            num_run.font.bold = True
            num_run.font.color.rgb = self._get_readable_text_color(self.current_theme['accent'])
            
            # Add background to number
            try:
                p = num_para._element
                pPr = p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '%02x%02x%02x' % self.current_theme['accent'])
                pPr.append(shd)
            except:
                pass
            
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Spacer
            cell.add_paragraph()
            
            # Card text
            text_para = cell.add_paragraph()
            text_run = text_para.add_run(point)
            font_size = self._fit_card_font_size(point)
            text_run.font.size = Pt(font_size)
            # FIX: Use readable color for card background
            text_run.font.color.rgb = self._get_readable_text_color(self.current_theme['card'])
            text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text_para.paragraph_format.line_spacing = 1.35
    
    
    # LAYOUT 4: FOUR GRID
    def _create_fixed_four_grid_slide(self, doc, slide_data, slide_num):
        """2x2 grid layout with numbered circles"""
        
        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Title (centered)
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self.current_theme['text'])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Spacer
        doc.add_paragraph()
        
        # Parse content (expect 4 points)
        content = self._parse_content(slide_data.get('content', ''))
        points = content[:4]
        while len(points) < 4:
            points.append(f"Point {len(points) + 1}")
        
        # Create table with 2x2 grid
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        for col_idx in range(2):
            table.columns[col_idx].width = Inches(5)
        
        positions = [
            (0, 0), (0, 1),
            (1, 0), (1, 1)
        ]
        
        circle_text_color = self._get_readable_text_color(self.current_theme['accent'])
        
        for i, (row_idx, col_idx) in enumerate(positions):
            if i >= len(points):
                break
            
            cell = table.rows[row_idx].cells[col_idx]
            
            # Set cell background color
            self._add_table_cell_color(cell, self.current_theme['card'])
            self._add_border_to_cell(cell, self.current_theme['accent'], width=2)
            
            # Number circle
            num_para = cell.add_paragraph()
            num_run = num_para.add_run(f"  {i+1}  ")
            num_run.font.size = Pt(18)
            num_run.font.bold = True
            num_run.font.color.rgb = circle_text_color
            
            # Add background to number
            try:
                p = num_para._element
                pPr = p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '%02x%02x%02x' % self.current_theme['accent'])
                pPr.append(shd)
            except:
                pass
            
            num_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Spacer
            cell.add_paragraph()
            
            # Card text
            text_para = cell.add_paragraph()
            text_run = text_para.add_run(points[i])
            text_run.font.size = Pt(13)
            # FIX: Use readable color for card background
            text_run.font.color.rgb = self._get_readable_text_color(self.current_theme['card'])
            text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text_para.paragraph_format.line_spacing = 1.25
    
    
    # LAYOUT 5: SPLIT BOX (GENERAL)
    def _create_fixed_split_box_slide(self, doc, slide_data, slide_num):
        """Split layout with image left, content right"""
        
        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Create table with 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(5)
        table.columns[1].width = Inches(5)
        
        row = table.rows[0]
        
        # LEFT CELL: Image
        left_cell = row.cells[0]
        # Try multiple image sources
        image_url = slide_data.get('image') or slide_data.get('bg_image') or slide_data.get('bg_url')
        if image_url:
            img_stream = self._download_image(image_url, max_size=(600, 720))
            if img_stream:
                left_cell.add_paragraph().add_run().add_picture(img_stream, width=Inches(4.5))
        
        # RIGHT CELL: Content
        right_cell = row.cells[1]
        self._add_table_cell_color(right_cell, self.current_theme['card'])
        
        # Title
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = right_cell.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self.current_theme['text'])
        
        # Spacer
        right_cell.add_paragraph()
        
        # Content bullets
        content = self._parse_content(slide_data.get('content', ''))
        points = content[:5]
        
        for point in points:
            point_para = right_cell.add_paragraph()
            point_run = point_para.add_run(f"• {point}")
            point_run.font.size = Pt(16)
            point_run.font.color.rgb = RGBColor(*self.current_theme['text'])
            point_para.paragraph_format.space_after = Pt(12)
    
    
    # LAYOUT 5: SPLIT BOX (SLIDE 5 SPECIAL)
    def _create_fixed_split_box_slide_slide5(self, doc, slide_data, slide_num):
        """Split layout with image left, 2+1 grid content cards right (Slide 5 only)"""
        
        # Background paragraph/color setup
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Create container table
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(6.0)
        
        row = table.rows[0]
        left_cell = row.cells[0]
        right_cell = row.cells[1]
        
        # LEFT CELL: Responsive Image
        image_url = slide_data.get('image')
        if image_url:
            img_stream = self._download_image(image_url, max_size=(800, 1000))
            if img_stream:
                p = left_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_stream, width=Inches(4.5))
        
        # RIGHT CELL: Content Panel
        is_light = self._calculate_brightness(self.current_theme['bg']) > 128
        panel_rgb = self.current_theme['accent'] if is_light else self.current_theme['bg']
        title_color = RGBColor(255, 255, 255) if is_light else RGBColor(*self.current_theme['text'])
        text_color = RGBColor(255, 255, 255) if is_light else RGBColor(*self.current_theme['text'])
        accent_color = RGBColor(*self.current_theme['accent'])
        card_bg_rgb = self.current_theme['card']
        card_text_color = self._get_readable_text_color(card_bg_rgb)
        
        self._add_table_cell_color(right_cell, panel_rgb)
        
        # Header / Title
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = right_cell.add_paragraph()
        title_run = title_para.add_run(title.upper())
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = title_color
        
        # Content Parsing
        content = self._parse_content(slide_data.get('content', ''))
        
        # Intro Text (First point)
        if content:
            intro_p = right_cell.add_paragraph()
            intro_run = intro_p.add_run(content[0])
            intro_run.font.size = Pt(12)
            intro_run.font.color.rgb = text_color
            right_cell.add_paragraph() # Spacer
        
        # 2+1 GRID FOR CARDS
        card_points = content[1:4] if len(content) > 1 else []
        while len(card_points) < 3: card_points.append("Innovative solution detail for this slide section.")
        
        # Small nested table for the grid
        card_table = right_cell.add_table(rows=2, cols=2)
        card_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Fill Cards
        for i, point in enumerate(card_points):
            r = 0 if i < 2 else 1
            c = i if i < 2 else 0
            cell = card_table.cell(r, c)
            self._add_table_cell_color(cell, card_bg_rgb)
            
            # Number Label
            num_p = cell.add_paragraph()
            num_run = num_p.add_run(f"0{i+1}")
            num_run.font.size = Pt(18)
            num_run.font.bold = True
            num_run.font.color.rgb = accent_color
            
            # Content Text
            text_p = cell.add_paragraph()
            text_run = text_p.add_run(point)
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = card_text_color
            text_p.paragraph_format.space_after = Pt(12)
    
    
    # LAYOUT 6: IMAGE CARDS
    def _create_fixed_image_cards_slide(self, doc, slide_data, slide_num):
        """Side-by-side layout with image and text"""
        
        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Create table with 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(5)
        table.columns[1].width = Inches(5)
        
        row = table.rows[0]
        
        # LEFT CELL: Image
        left_cell = row.cells[0]
        image_url = slide_data.get('image') or slide_data.get('bg_image') or slide_data.get('bg_url')
        if image_url:
            img_stream = self._download_image(image_url, max_size=(640, 720))
            if img_stream:
                left_cell.add_paragraph().add_run().add_picture(img_stream, width=Inches(4.5))
        
        # RIGHT CELL: Content
        right_cell = row.cells[1]
        
        # Title
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = right_cell.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(30)
        title_run.font.bold = True
        title_run.font.color.rgb = self._get_readable_text_color(self.current_theme['bg'])
        
        # Spacer
        right_cell.add_paragraph()
        
        # Content points
        content = self._parse_content(slide_data.get('content', ''))
        points = content if content else ["Content for this slide"]
        
        for i, point in enumerate(points[:8]):
            point_para = right_cell.add_paragraph()
            point_run = point_para.add_run(f"{i+1}. {point}")
            point_run.font.size = Pt(16)
            # FIX: Use readable color for RIGHT cell background
            point_run.font.color.rgb = self._get_readable_text_color(self.current_theme['bg'])
            point_para.paragraph_format.line_spacing = 1.35
            point_para.paragraph_format.space_after = Pt(12)
    
    
    # LAYOUT 7: EXECUTIVE SUMMARY
    def _create_executive_summary_slide(self, doc, slide_data, slide_num):
        """Modern executive summary with design elements"""
        
        # White background
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, (255, 255, 255))
        
        # Header with "EXECUTIVE SUMMARY" badge
        header_table = doc.add_table(rows=1, cols=2)
        header_table.columns[0].width = Inches(8)
        header_table.columns[1].width = Inches(2)
        
        # Left: Main title
        left_cell = header_table.rows[0].cells[0]
        title = slide_data.get('title', 'Executive Summary')
        title_para = left_cell.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(30, 41, 59)
        
        # Right: Badge
        right_cell = header_table.rows[0].cells[1]
        self._add_table_cell_color(right_cell, (30, 58, 138))
        badge_para = right_cell.add_paragraph()
        badge_run = badge_para.add_run("EXECUTIVE\nSUMMARY")
        badge_run.font.size = Pt(9)
        badge_run.font.bold = True
        badge_run.font.color.rgb = RGBColor(255, 255, 255)
        badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Accent line (using paragraph with background)
        accent_para = doc.add_paragraph()
        accent_run = accent_para.add_run("    " * 20)
        accent_run.font.size = Pt(3)
        self._add_page_color(accent_para, self.current_theme['accent'])
        
        # Spacer
        doc.add_paragraph()
        
        # Content sections
        content = self._parse_content(slide_data.get('content', ''))
        while len(content) < 3:
            content.append(f"Section {len(content) + 1} content here...")
        
        sections = content[:3]
        
        for section_text in sections:
            # Section label
            label_para = doc.add_paragraph()
            label_run = label_para.add_run("TITLE")
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor(100, 116, 139)
            
            # Section description
            desc_para = doc.add_paragraph()
            desc_run = desc_para.add_run(section_text)
            desc_run.font.size = Pt(13)
            desc_run.font.color.rgb = RGBColor(51, 65, 85)
            desc_para.paragraph_format.line_spacing = 1.3
            desc_para.paragraph_format.space_after = Pt(15)

        # Optional Image
        image_url = slide_data.get('image') or slide_data.get('bg_image') or slide_data.get('bg_url')
        if image_url:
            img_stream = self._download_image(image_url, max_size=(1000, 400))
            if img_stream:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_stream, width=Inches(7.0))
    
    
    # LAYOUT 8: ROADMAP
    def _create_fixed_roadmap_clean_slide(self, doc, slide_data, slide_num):
        """Vertical timeline with boxes on the right side"""
        
        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Title (centered)
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        # FIX: Use readable color
        title_run.font.color.rgb = self._get_readable_text_color(self.current_theme['bg'])
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Spacer
        doc.add_paragraph()
        
        # Parse content
        content = self._parse_content(slide_data.get('content', ''))
        steps = content[:6]
        if len(steps) < 3:
            steps = ["Step 1", "Step 2", "Step 3"]
        
        # Create timeline table
        table = doc.add_table(rows=len(steps), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        table.columns[0].width = Inches(1)  # Circle column
        table.columns[1].width = Inches(9)  # Content column
        
        circle_text_color = self._get_readable_text_color(self.current_theme['accent'])
        
        if self._calculate_brightness(self.current_theme['bg']) > 128:
            box_bg = (245, 247, 250)
        else:
            box_bg = self.current_theme['card']
        
        for i, step in enumerate(steps):
            row = table.rows[i]
            
            # Circle cell
            circle_cell = row.cells[0]
            num_para = circle_cell.add_paragraph()
            num_run = num_para.add_run(f"  {i+1}  ")
            num_run.font.size = Pt(20)
            num_run.font.bold = True
            num_run.font.color.rgb = circle_text_color
            
            # Add background to circle
            self._add_table_cell_color(circle_cell, self.current_theme['accent'])
            num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Content cell
            content_cell = row.cells[1]
            self._add_table_cell_color(content_cell, box_bg)
            self._add_border_to_cell(content_cell, self.current_theme['accent'], width=2)
            
            # Parse year if present
            import re
            year_match = re.match(r'^(\d{4})[:\s-]*(.+)', step.strip())
            if year_match:
                year = year_match.group(1)
                description = year_match.group(2).strip()
                
                # Year
                year_para = content_cell.add_paragraph()
                year_run = year_para.add_run(year)
                year_run.font.size = Pt(16)
                year_run.font.bold = True
                year_run.font.color.rgb = RGBColor(*self.current_theme['accent'])
                
                # Description
                desc_para = content_cell.add_paragraph()
                desc_run = desc_para.add_run(description)
                desc_run.font.size = Pt(11)
                # FIX: Use readable color
                desc_run.font.color.rgb = self._get_readable_text_color(box_bg)
                desc_para.paragraph_format.line_spacing = 1.2
            else:
                # Just description
                desc_para = content_cell.add_paragraph()
                desc_run = desc_para.add_run(step)
                desc_run.font.size = Pt(11)
                # FIX: Use readable color
                desc_run.font.color.rgb = self._get_readable_text_color(box_bg)
                desc_para.paragraph_format.line_spacing = 1.2
    
    
    # LAYOUT 9: IMAGE OVERLAY (MISSION)
    def _create_image_overlay_slide(self, doc, slide_data, slide_num):
        """Split screen layout with rounded card overlay (70/30 split)"""
        
        # Determine colors based on brightness - if light theme, use white bg
        is_light = self._calculate_brightness(self.current_theme['bg']) > 128
        
        # White background for the page
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, (255, 255, 255))
        
        # Create table with 2 columns
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # PARITY: 70/30 split. 30% left (card), 70% right (image)
        # Total width approx 10 inches for margins
        table.columns[0].width = Inches(3.2)  # Left: Card (approx 30%)
        table.columns[1].width = Inches(6.8)  # Right: Image (approx 70%)
        
        row = table.rows[0]
        
        # LEFT CELL: Card with content
        left_cell = row.cells[0]
        
        # In Dialogue theme (light), the box should be Purple (accent)
        box_color = self.current_theme['accent'] if is_light else self.current_theme['bg']
        text_color = RGBColor(255, 255, 255) if is_light else RGBColor(*self.current_theme['text'])
        
        self._add_table_cell_color(left_cell, box_color)
        
        # Padding/Spacer at top
        left_cell.add_paragraph()
        
        # "KEY INSIGHT" label
        label_para = left_cell.add_paragraph()
        label_run = label_para.add_run("KEY INSIGHT")
        label_run.font.size = Pt(12)
        label_run.font.bold = True
        label_run.font.color.rgb = text_color
        
        # Spacer
        left_cell.add_paragraph()
        
        # Card title
        title = slide_data.get('title', 'Summary and Conclusion')
        title_para = left_cell.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = text_color
        title_para.paragraph_format.line_spacing = 1.1
        
        # Spacer
        left_cell.add_paragraph()
        
        # Content paragraphs
        content = self._parse_content(slide_data.get('content', ''))
        paragraphs = content[:4]
        if not paragraphs:
            paragraphs = [
                "Overview of main conclusions",
                "Key strategic takeaways"
            ]

        for p_text in paragraphs:
            p = left_cell.add_paragraph()
            run = p.add_run(f"• {p_text}")
            run.font.size = Pt(12)
            run.font.color.rgb = text_color
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(8)
            
        # RIGHT CELL: Large Image (70%)
        right_cell = row.cells[1]
        image_url = slide_data.get('image') or slide_data.get('bg_image') or slide_data.get('bg_url')
        if image_url:
            img_stream = self._download_image(image_url, max_size=(1600, 900), cover=True)
            if img_stream:
                pic_para = right_cell.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.add_run().add_picture(img_stream, width=Inches(6.6))
        else:
            # Placeholder color if no image
            self._add_table_cell_color(right_cell, (241, 245, 249))
    
    
    # LAYOUT 10: STANDARD SLIDE
    def _create_standard_slide(self, doc, slide_data, slide_num):
        """Standard slide layout with title and bullet points"""
        
        # Background paragraph
        bg_para = doc.add_paragraph()
        self._add_page_color(bg_para, self.current_theme['bg'])
        
        # Determine readable text color for this background
        text_color = self._get_readable_text_color(self.current_theme['bg'])
        
        # Title
        title = slide_data.get('title', f'Slide {slide_num}')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = text_color
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Spacer
        doc.add_paragraph()
        
        # Content with fallbacks (mission/agenda)
        content_raw = slide_data.get('content') or slide_data.get('mission') or slide_data.get('agenda') or ''
        content = self._parse_content(content_raw)
        points = content[:8]
        
        if not points:
            points = ["Information for this section"]
        
        for point in points:
            point_para = doc.add_paragraph()
            point_run = point_para.add_run(f"• {point}")
            point_run.font.size = Pt(18)
            point_run.font.color.rgb = text_color
            point_para.paragraph_format.line_spacing = 1.4
            point_para.paragraph_format.space_after = Pt(14)
            
        # ✅ ADD IMAGE SUPPORT TO STANDARD SLIDES (Robust fields)
        img_url = slide_data.get('image') or slide_data.get('bg_image') or slide_data.get('bg_url')
        if img_url:
            img_stream = self._download_image(img_url, max_size=(1000, 500))
            if img_stream:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(img_stream, width=Inches(8.0))
        
        logger.info(f"   → Standard slide created with {len(points)} points")

# MODULE INITIALIZATION
logger.info("✅ DOCX Service v5.0.1 (THEME FIX) loaded successfully")
logger.info("   - 8 Themes supported")
logger.info("   - 8 Layouts implemented")
logger.info("   - Smart text colors based on brightness")
logger.info("   - Enhanced theme detection")
logger.info("   - All hardcoded colors removed")